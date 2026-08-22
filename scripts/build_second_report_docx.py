from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "report" / "entwuerfe" / "2_Zwischenstand_Jann_Koerner_bearbeitbar.docx"
DATA = json.loads((ROOT / "public" / "data" / "energy.json").read_text(encoding="utf-8"))
FIGURES = ROOT / "experiments" / "figures"

NAVY = "21344D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6472"
LIGHT = "F4F6F9"
GOLD = "C79A22"
RED = "B84040"
GREEN = "4F8A4A"
WHITE = "FFFFFF"
BLACK = "20242A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = round(sum(widths_in) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = round(widths_in[index] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, italic=None, color=BLACK, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text: str, style=None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix) :])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_font(run)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
        for index, value in enumerate(values):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, size=8.6)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title: str, body: str, accent: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F9FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=9.5)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, filename: str, width_cm: float, caption: str, alt_text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(FIGURES / filename), width=Cm(width_cm))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.keep_with_next = False
    r = cap.add_run(caption)
    set_font(r, size=9, italic=True, color=MUTED)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=8, color=MUTED)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Informationsvisualisierung · Zweiter Zwischenstand")
    set_font(r, size=8, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "2. Zwischenstand: Deutschlands Rolle im europäischen Stromnetz"
    doc.core_properties.author = "Jann Körner"
    doc.core_properties.subject = "Elm-Implementierung und Experimente"

    # Editorial cover, deliberately A4 as required by the university context.
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("ZWEITER ZWISCHENSTAND")
    set_font(r, size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Deutschlands Rolle im\neuropäischen Stromnetz")
    set_font(r, size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("Beschreibung der Elm-Implementierung und Experimente")
    set_font(r, size=14, italic=True, color=MUTED)

    for label, value in (
        ("Autor", "Jann Körner"),
        ("Modul", "Informationsvisualisierung"),
        ("Dokument", "Bearbeitbarer zweiter Zwischenstand"),
        ("Stand", "21. August 2026"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        a = p.add_run(f"{label}: ")
        set_font(a, size=11, bold=True, color=NAVY)
        b = p.add_run(value)
        set_font(b, size=11, color=BLACK)

    doc.add_page_break()
    add_heading(doc, "Stand dieses Zwischenstands", 1)
    add_text(doc, "Dieser Zwischenstand dokumentiert Screenshots einer kompilierten Elm-Anwendung. Die drei SVG-Ansichten sind über ein gemeinsames Model interaktiv verbunden. Auswahländerungen in einer Ansicht aktualisieren die beiden anderen Ansichten.")
    add_callout(
        doc,
        "Datenbankstatus",
        "Der Datensatz wurde über die ScienceData-PostgREST-Schnittstelle aus dem Schema energycharts exportiert. Verwendet werden die Views v_cbpf, v_cbet, v_price und v_totalpower. Alle Abfragen sind auf Deutschland beziehungsweise DE-LU, den Zeitraum 01.-02.01.2025 und feste Zeilenlimits beschränkt. Passwort und Token sind nicht Bestandteil des Repositories.",
        GREEN,
    )
    add_heading(doc, "Gliederung", 2)
    for item in (
        "1 Einleitung",
        "2 Daten",
        "3 Visualisierungen",
        "4 Implementierung",
        "5 Anwendungsfälle",
        "6 Verwandte Arbeiten",
        "7 Zusammenfassung und Ausblick",
        "Anhang: Git-Historie und KI-Unterstützung",
        "Literatur",
    ):
        add_bullet(doc, item)

    add_heading(doc, "1 Einleitung", 1)
    add_heading(doc, "1.1 Anwendungshintergrund", 2)
    add_text(doc, "Der europäische Stromverbund gleicht Erzeugung und Verbrauch über Ländergrenzen hinweg aus. Ein Land kann innerhalb eines Tages mehrfach zwischen Import und Export wechseln. Einzelne Tabellenabfragen liefern Maxima oder Summen, zeigen aber nicht unmittelbar, ob ein Ereignis Teil eines räumlich-zeitlichen Musters ist. Die Anwendung verbindet deshalb Netzstruktur, zeitlichen Verlauf und verdichtete Musterübersicht.")
    add_text(doc, "Die Forschungsfrage lautet: Welche stabilen oder wiederkehrenden Muster zeigen die physischen grenzüberschreitenden Stromflüsse europäischer Länder, und wie fallen diese Muster zeitlich mit Erzeugungsmix und Strompreisen zusammen? Deutschland ist der primäre Fokus; die Architektur soll auf weitere freigegebene europäische Länder übertragbar bleiben.")
    add_heading(doc, "1.2 Zielgruppen", 2)
    add_text(doc, "Primäre Zielgruppe sind Datenjournalist:innen in deutschsprachigen Energie- und Klimaredaktionen. Vorausgesetzt werden Grundkenntnisse zu europäischen Ländern, Leistung in GW, Energie in GWh und Zeitreihen. SQL- oder API-Kenntnisse sind nicht erforderlich. Die Anwendung soll dominante Austauschbeziehungen, Richtungswechsel und auffällige Perioden sichtbar machen und mit dem zeitgleichen Erzeugungsmix kontextualisieren.")
    add_heading(doc, "1.3 Überblick und Beiträge", 2)
    add_text(doc, "Der aktuelle Beitrag besteht aus einer kompilierten Elm-Anwendung mit drei verbundenen Visualisierungen, einem reproduzierbaren PostgreSQL/PostgREST-Export und dokumentierten Interaktionstests. Die Implementierung trennt Domänentypen, JSON-Decodierung, Anwendungszustand und SVG-Views.")

    add_heading(doc, "2 Daten", 1)
    add_text(doc, "Die Primärquelle ist die von der Veranstaltung bereitgestellte PostgreSQL-Datenbank. Sie wird über die ScienceData-PostgREST-Schnittstelle angesprochen. Das Exportskript fragt die benötigten Views mit Filtern und Limits ab, normalisiert die Ergebnisse und stellt sie als statische JSON-Datei bereit, die Elm anschließend per HTTP lädt. Zugangsdaten werden nicht in Elm oder Git gespeichert.")
    add_table(
        doc,
        ["Funktion", "Endpunkt", "Prüfergebnis"],
        [
            ["Authentifizierung", "/sciencedata/token", "HTTP 200; Token erhalten"],
            ["Schemaauswahl", "Accept-Profile: energycharts", "66 Ressourcen sichtbar"],
            ["Flüsse und Handel", "v_cbpf, v_cbet", "je 2304 Zeilen, begrenzt"],
            ["Erzeugung und Preis", "v_totalpower, v_price", "192 bzw. 48 Zeilen"],
        ],
        [1.55, 2.55, 2.4],
    )
    add_text(doc, "Für die Analyse werden physische Flüsse aus v_cbpf, Handelswerte aus v_cbet, DE-LU-Preise aus v_price und deutsche Erzeugungswerte aus v_totalpower kombiniert. v_publicpower wurde geprüft, enthält für Deutschland jedoch Nullwerte. v_totalpower ist gefüllt; seine deutschen Erzeugungswerte werden von MW in GW umgerechnet.")
    add_heading(doc, "Verwendeter Datenausschnitt", 2)
    add_text(doc, f"Für die Implementierungs- und Interaktionstests lädt Elm {len(DATA['samples'])} Stunden des Zeitraums {DATA['period']} aus public/data/energy.json. Die Datei wurde mit scripts/build_postgrest_fixture.py direkt aus den vier PostgreSQL-Views erzeugt. Sie enthält elf Partnerländer. Viertelstundenwerte wurden auf Stundenmittel aggregiert, damit sie zu den Stundenpreisen passen.")
    add_callout(doc, "Semantische Grenze", "Der Erzeugungsmix eines Partnerlands beschreibt die zeitgleiche Produktion. Er weist nicht nach, aus welchen Energieträgern eine konkrete importierte Strommenge stammt. Die Anwendung verwendet deshalb keine Formulierungen wie „importierter Windstrom“.", BLUE)

    add_heading(doc, "3 Visualisierungen", 1)
    add_heading(doc, "3.1 Analyse der Anwendungsaufgaben", 2)
    add_table(
        doc,
        ["ID", "Aufgabe", "Visueller Mehrwert"],
        [
            ["A1", "dominante gerichtete Länderbeziehungen erkennen", "viele Relationen und beide Richtungen gleichzeitig vergleichen"],
            ["A2", "Richtungswechsel und wiederkehrende Muster finden", "Form und Wiederholung statt nur Einzelwerte erfassen"],
            ["A3", "Flussperioden mit Erzeugungsmix abgleichen", "synchronisierte Zeitachsen und gemeinsame Auswahl"],
            ["A4", "vom Überblick in Paar und Zeitpunkt wechseln", "iterative Exploration über gekoppelte Ansichten"],
        ],
        [0.45, 2.75, 3.3],
    )
    add_heading(doc, "3.2 Anforderungen an die Visualisierungen", 2)
    add_text(doc, "Richtung und Betrag werden redundant codiert. Farbe unterscheidet Import und Export; Pfeilspitzen zeigen die Richtung; Linienbreite beziehungsweise Farbintensität zeigt den Betrag. Alle Ansichten verwenden dieselbe Vorzeichenkonvention. Fehlende Werte dürfen nicht als Null erscheinen. Auswahl und Datenstatus müssen sichtbar sein.")
    add_heading(doc, "3.3 Präsentation der Visualisierungen", 2)
    add_heading(doc, "3.3.1 Visualisierung Eins - gerichtete Flussansicht", 3)
    add_figure(doc, "elm_chord.png", 11.0, "Abbildung 1: Gerichtete Flussansicht, von Elm als SVG gerendert. Frankreich ist ausgewählt; übrige Verbindungen werden abgeblendet.", "Elm-SVG der gerichteten Stromflüsse zwischen Deutschland und Partnerländern")
    add_text(doc, "Deutschland und die Partnerländer werden radial angeordnet. Die gerichteten Verbindungen zeigen die physischen Flüsse des ausgewählten Zeitpunkts. Die Ansicht ist ein fokussierter Chord-/Flussprototyp; für eine vollständige europäische Chord-Matrix sind weitere freigegebene Länder-zu-Länder-Daten erforderlich.")
    add_heading(doc, "3.3.2 Visualisierung Zwei - gestapelte Zeitreihe", 3)
    add_figure(doc, "elm_timeline.png", 15.2, "Abbildung 2: Überarbeitete Erzeugungszeitreihe aus Elm. Die obere Y-Achse zeigt die absolute Erzeugungsleistung in GW. Der physische Fluss zu Frankreich wird darunter mit eigener symmetrischer GW-Achse dargestellt; die gestrichelte Linie markiert den gewählten Zeitpunkt.", "Elm-SVG der absoluten Erzeugungsleistung mit separatem Stromflussdiagramm")
    add_text(doc, "Die gestapelten Flächen zeigen absolute Leistungen für Erneuerbare, Kohle, Gas und Sonstige. Teilstriche und Achsenbeschriftung machen die Größenordnung in GW ablesbar. Unterhalb des Erzeugungsmixes besitzt der physische Stromfluss des ausgewählten Partnerlands ein separates Diagramm mit eigener symmetrischer Skala um null. Dadurch werden Stromfluss und Erzeugungsleistung nicht fälschlich auf derselben Y-Skala verglichen.")
    add_text(doc, "Für die ausgewählte Stunde zeigt die Anwendung Gesamtleistung, absolute Werte und prozentuale Anteile der vier Erzeugungsgruppen. Zusätzlich werden Strompreis und physischer Fluss des gewählten Partnerlands numerisch ausgegeben. Damit unterstützt die Ansicht sowohl die Mustererkennung im Verlauf als auch das Ablesen konkreter Einzelwerte.")
    add_heading(doc, "3.3.3 Visualisierung Drei - Pixelmatrix", 3)
    add_figure(doc, "elm_matrix.png", 16.0, "Abbildung 3: Pixelmatrix aus Elm. Zeilen sind Partnerländer, Spalten Stunden; Rot bedeutet Import, Blau Export.", "Elm-SVG einer Pixelmatrix der Stromflüsse nach Partnerland und Stunde")
    add_text(doc, "Die Matrix zeigt stabil gerichtete und wechselnde Beziehungen kompakt. Die schwarz umrandete Zelle markiert die gemeinsame Auswahl. Für längere Zeiträume sind Zoom, Aggregation und eine explizite Codierung fehlender Werte vorgesehen.")
    add_heading(doc, "3.4 Interaktion", 2)
    add_text(doc, "Ein Klick auf ein Partnerland oder eine Verbindung setzt selectedPartner. Ein Klick in die Zeitreihe setzt selectedIndex. Ein Klick auf eine Matrixzelle setzt beide Werte gleichzeitig. Alle Views werden anschließend aus demselben Model neu gerendert. Die Schaltfläche „Auswahl zurücksetzen“ entfernt den Paarfilter und springt zum ersten Zeitpunkt.")

    add_heading(doc, "4 Implementierung", 1)
    add_text(doc, "Der Prototyp wurde mit Elm 0.19.1 implementiert. elm make kompiliert den Stand erfolgreich nach public/elm.js. Die Anwendung lädt ihren Datensatz über Http.get. Darstellung und Interaktion werden ausschließlich durch Elm und SVG erzeugt; Python wird nicht für die Diagrammgeometrie verwendet.")
    add_table(
        doc,
        ["Modul", "Verantwortung"],
        [
            ["src/Domain.elm", "Datentypen Dataset, Sample, Generation und Flow"],
            ["src/Api.elm", "HTTP-Anfrage und JSON-Decoder"],
            ["src/Main.elm", "Model, Msg, update und Seitenaufbau"],
            ["src/View/Chord.elm", "gerichtete SVG-Pfade und Partnerauswahl"],
            ["src/View/TimeSeries.elm", "gestapelte Flächen, Flusslinie und Zeitauswahl"],
            ["src/View/FlowMatrix.elm", "Pixelzellen, divergierende Skala und kombinierte Auswahl"],
        ],
        [2.1, 4.4],
    )
    add_text(doc, "Der Flow-Datentyp speichert den physischen Wert aus v_cbpf und den Handelswert aus v_cbet getrennt. Die gerichteten Kanten und die Pixelmatrix codieren den physischen Fluss. Ein SVG-Tooltip der Flussansicht zeigt zusätzlich beide Zahlenwerte für das gewählte Länderpaar.")
    add_heading(doc, "Gemeinsamer Zustand", 2)
    add_text(doc, "Der relevante Zustand besteht aus Dataset, selectedIndex und selectedPartner. SelectPartner ändert nur das Partnerland, SelectTime nur den Zeitpunkt, SelectCell beide Werte. Diese Nachrichten werden zentral in update verarbeitet. Die Views erhalten Daten und Callback-Nachrichten als Parameter; sie führen keine eigenen HTTP-Anfragen durch.")
    add_heading(doc, "Datenpipeline", 2)
    for step in (
        "Freigegebene Tabellen, Views und Spalten über die PostgREST-OpenAPI ermitteln.",
        "Daten mit PostgREST-Filtern, Limits und Offsets seitenweise exportieren und normalisieren.",
        "JSON ohne Token unter public/data bereitstellen.",
        "Datensatz durch Api.elm über HTTP laden und decodieren.",
        "Auswahlzustand in Main.elm aktualisieren und alle SVG-Views neu berechnen.",
    ):
        add_number(doc, step)
    add_text(doc, "Der Datenpfad wurde vollständig ausgeführt. Der Export erzeugte aus 4848 begrenzt abgefragten Rohzeilen 48 normalisierte Stunden mit elf Partnerländern. Der Datenstatus im Kopf der Elm-Anwendung nennt das Schema und die verwendeten Views.")

    add_heading(doc, "5 Anwendungsfälle", 1)
    add_text(doc, "Die folgenden Experimente prüfen den technischen Stand und die Kopplung der Ansichten. Sie sind keine energiewirtschaftliche Evaluation. Der kurze Zweitageszeitraum dient nur als reproduzierbarer Testdatensatz.")
    add_heading(doc, "5.1 Anwendung Visualisierung Eins", 2)
    add_text(doc, "Im Browser wurde Frankreich in der gerichteten Flussansicht ausgewählt. Die Toolbar wechselte zu „Auswahl: France“, die übrigen Verbindungen wurden abgeblendet und in der Zeitreihe erschien die violette Linie des Länderpaars. Damit wirkt die Selektion über die Grenzen der ersten View hinaus.")
    add_heading(doc, "5.2 Anwendung Visualisierung Zwei", 2)
    add_text(doc, "Anschließend wurde in der Zeitreihe der Zeitpunkt 02.01. 00 h angeklickt. Dieser Zeitpunkt besitzt im nullbasierten Datensatz den selectedIndex 24. France blieb als Partnerauswahl erhalten. Die gestrichelte Markierung, die Chord-Werte und die Matrixauswahl wurden aus demselben Index abgeleitet. Die Detailanzeige wies für diesen Zeitpunkt 60,2 GW Gesamtleistung sowie die absoluten und prozentualen Beiträge der Erzeugungsgruppen aus; der Frankreich-Fluss betrug +0,2 GW.")
    add_heading(doc, "5.3 Anwendung Visualisierung Drei", 2)
    add_text(doc, "Ein Klick auf eine Zelle der Pixelmatrix setzte Partnerland und Stunde gemeinsam. Im Test zeigte die Toolbar anschließend „Denmark · 01.01. 12 h“. Die Matrix erfüllt damit ihre Rolle als Überblick und direkter Einstieg in einen Detailzustand.")
    add_table(
        doc,
        ["Prüfung", "Ergebnis"],
        [
            ["Elm-Build", "erfolgreich; sechs Module kompiliert"],
            ["HTTP-Laden", "48 Stunden aus public/data/energy.json"],
            ["Chord-Auswahl", "Partnerfilter in allen Ansichten sichtbar"],
            ["Zeitwahl", "selectedIndex, GW-Detailwerte und Anteilsausgabe aktualisiert"],
            ["Matrixzelle", "Partner und Zeitpunkt gemeinsam aktualisiert"],
            ["PostgreSQL", "vier Energy-Charts-Views erfolgreich exportiert"],
        ],
        [2.3, 4.2],
    )

    add_heading(doc, "6 Verwandte Arbeiten", 1)
    add_text(doc, "Munzners Nested Model trennt Domänenproblem, Daten- und Aufgabenabstraktion, visuelle Codierung sowie Algorithmus [1]. Diese Ebenen strukturieren die Begründung der drei Ansichten. Roberts beschreibt koordinierte Multiple Views als gemeinsames Explorationssystem [2]; selectedPartner und selectedIndex realisieren hier Brushing and Linking.")
    add_text(doc, "Keim systematisiert pixelorientierte Verfahren für große Datenmengen [3]. Die Matrix ordnet Pixel fachlich als Partnerland × Zeit. Aigner et al. behandeln Entwurfsentscheidungen für zeitbezogene Daten und begründen die explizite Wahl von Granularität und synchronisierten Zeitachsen [4]. Das Directed-Chord-Beispiel von Observable/D3 dient nur als Gestaltungsreferenz, nicht als wissenschaftliche Evaluation [6].")

    add_heading(doc, "7 Zusammenfassung und Ausblick", 1)
    add_text(doc, "Der zweite Zwischenstand enthält nun eine reale, kompilierte Elm-Implementierung der drei verbundenen Ansichten. Die Abbildungen stammen aus der laufenden Elm-Anwendung. Datenladung, Decoder, Auswahlzustand und SVG-Views sind getrennt dokumentiert.")
    add_text(doc, "Der PostgreSQL-Export ist für einen begrenzten Zweitagesausschnitt umgesetzt und die Elm-Screenshots wurden mit diesem Datensatz neu erzeugt. Offen bleiben längere Untersuchungszeiträume, Zoom und Aggregation, ein Top-k-Filter sowie ein kleiner Nutzertest. Zusätzlich sollte die ungewöhnliche Einheitendarstellung von v_totalpower vor der Endfassung mit der Datenbankdokumentation abgeglichen werden.")

    add_heading(doc, "Anhang: Git-Historie", 1)
    add_text(doc, "Das Repository enthält Bericht, Elm-Quellcode, Builddateien, Datenprüfung, Experimente und Markdown-Dokumentation. Die Endfassung ergänzt an dieser Stelle eine exportierte Git-Historie mit nachvollziehbaren Commits für Datenzugriff, jede View, Interaktion, Tests und Bericht.")
    add_heading(doc, "KI-Unterstützung", 2)
    add_text(doc, "Bei Strukturierung, Formulierung und Codeerstellung wurde ein KI-Werkzeug eingesetzt. Die übernommenen Teile wurden durch elm make, Browserinteraktionen und visuelle Prüfung kontrolliert. Datenbankergebnisse werden nur dokumentiert, wenn sie tatsächlich abgerufen wurden. Vor der Einreichung muss der Autor Quellcode, Datenwerte, Literatur und Formulierungen eigenständig prüfen und die hochschulischen Kennzeichnungsregeln beachten.")

    add_heading(doc, "Literatur", 1)
    references = [
        "[1] Munzner, T. (2009). A Nested Model for Visualization Design and Validation. IEEE Transactions on Visualization and Computer Graphics, 15(6), 921-928. https://doi.org/10.1109/TVCG.2009.111",
        "[2] Roberts, J. C. (2007). State of the Art: Coordinated & Multiple Views in Exploratory Visualization. Proceedings of CMV 2007, 61-71. https://doi.org/10.1109/CMV.2007.20",
        "[3] Keim, D. A. (2000). Designing Pixel-Oriented Visualization Techniques: Theory and Applications. IEEE Transactions on Visualization and Computer Graphics, 6(1), 59-78. https://doi.org/10.1109/2945.841121",
        "[4] Aigner, W., Miksch, S., Schumann, H., & Tominski, C. (2011). Visualization of Time-Oriented Data. Springer. https://doi.org/10.1007/978-0-85729-079-3",
        "[5] Fraunhofer ISE. Energy-Charts API. https://api.energy-charts.info/ (abgerufen am 21.08.2026).",
        "[6] Observable/D3. Directed Chord Diagram. https://observablehq.com/@d3/directed-chord-diagram (abgerufen am 21.08.2026).",
        "[7] Martin-Luther-Universität Halle-Wittenberg. ScienceData-PostgREST-Schnittstelle, Schema energycharts. https://dbs.informatik.uni-halle.de/sciencedata/ (geprüft am 21.08.2026).",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(reference)
        set_font(r, size=9)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
